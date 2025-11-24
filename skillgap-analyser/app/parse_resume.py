# import os
# import re
# import json
# import numpy as np
# from PyPDF2 import PdfReader
# from sentence_transformers import SentenceTransformer
# import faiss
# import spacy

# # === Config ===
# threshold = 0.55   # similarity cutoff (higher = cleaner matches)
# debug = False      # set True to see top matches for every token

# # === Load FAISS + Metadata ===
# index_path = "vector_store/esco_faiss.index"
# meta_path = "vector_store/esco_meta.json"

# index = faiss.read_index(index_path)

# with open(meta_path, "r", encoding="utf-8") as f:
#     meta = json.load(f)

# # Detect format automatically
# if isinstance(meta, dict) and "skills" in meta:
#     print("📂 Detected format: dict with 'skills' key")
#     id2skill = {int(s["id"]): s for s in meta["skills"] if isinstance(s, dict)}
# elif isinstance(meta, dict):
#     print("📂 Detected format: dict with numeric/string keys")
#     id2skill = {int(k): v for k, v in meta.items() if k.isdigit()}
# elif isinstance(meta, list):
#     print("📂 Detected format: list")
#     id2skill = {int(s["id"]): s for s in meta if isinstance(s, dict) and "id" in s}
# else:
#     raise ValueError("❌ Unsupported meta format")

# print(f"✅ Loaded {len(id2skill)} ESCO skills into memory")

# # === Models ===
# model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
# try:
#     nlp = spacy.load("en_core_web_sm")  # lightweight for tokenization
# except OSError:
#     raise RuntimeError("⚠️ Run: python -m spacy download en_core_web_sm")


# # === Helpers ===
# def extract_text_from_pdf(pdf_path):
#     reader = PdfReader(pdf_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text


# def tokenize_text(text):
#     """Use spaCy for better tokenization, filter out junk and multi-line tokens."""
#     doc = nlp(text)
#     tokens = []
#     for token in doc:
#         t = token.text.strip().lower()
#         # clean weird line breaks and punctuation
#         t = re.sub(r"[\n\r,;]+", " ", t).strip()
#         if not t or len(t) < 2:
#             continue
#         if len(t.split()) > 3:  # skip long multi-word junk
#             continue
#         if not re.match(r"^[a-z0-9\-\+#\.]+$", t):
#             continue
#         tokens.append(t)
#     return list(set(tokens))


# def normalize_skills(tokens):
#     results = {}
#     embeddings = model.encode(tokens, convert_to_numpy=True, normalize_embeddings=True)

#     D, I = index.search(embeddings, 3)  # top-3 ESCO matches
#     for token, scores, ids in zip(tokens, D, I):
#         best_id = int(ids[0])
#         best_score = float(scores[0])
#         skill_data = id2skill.get(best_id, {})

#         if debug:
#             print(f"\n🔎 Token: {token}")
#             for idx, sc in zip(ids, scores):
#                 sd = id2skill.get(int(idx), {})
#                 print(f"   -> {sc:.3f} | {sd.get('preferredLabel')} ({sd.get('uri')})")

#         if best_score >= threshold and skill_data:
#             esco_skill = skill_data.get("preferredLabel")
#             uri = skill_data.get("uri")
#             # Deduplicate: keep highest score for same token/skill
#             key = (token, esco_skill)
#             if key not in results or best_score > results[key]["score"]:
#                 results[key] = {
#                     "token": token,
#                     "score": best_score,
#                     "esco_skill": esco_skill,
#                     "uri": uri,
#                 }

#     return list(results.values())


# def main(pdf_path):
#     text = extract_text_from_pdf(pdf_path)
#     tokens = tokenize_text(text)

#     print(f"\n📑 Extracted {len(tokens)} clean tokens from resume")

#     skills = normalize_skills(tokens)

#     output = {
#         "resume_name": os.path.basename(pdf_path),
#         "skills": skills,
#     }

#     out_path = os.path.splitext(pdf_path)[0] + "_skills.json"
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(output, f, indent=2)

#     print(f"\n✅ Results saved to {out_path}")


# if __name__ == "__main__":
#     import sys
#     pdf_path = sys.argv[1]
#     main(pdf_path)

import os
# import re
# import json
# import numpy as np
# from PyPDF2 import PdfReader
# from sentence_transformers import SentenceTransformer
# import faiss
# import spacy
# from collections import defaultdict

# # === Config ===
# threshold = 0.55   # similarity cutoff (higher = cleaner matches)
# debug = False      # set True to see top matches for every token

# # === Load FAISS + Metadata ===
# index_path = "vector_store/esco_faiss.index"
# meta_path = "vector_store/esco_meta.json"

# index = faiss.read_index(index_path)

# with open(meta_path, "r", encoding="utf-8") as f:
#     meta = json.load(f)

# # Detect format automatically
# if isinstance(meta, dict) and "skills" in meta:
#     print("📂 Detected format: dict with 'skills' key")
#     id2skill = {int(s["id"]): s for s in meta["skills"] if isinstance(s, dict)}
# elif isinstance(meta, dict):
#     print("📂 Detected format: dict with numeric/string keys")
#     id2skill = {int(k): v for k, v in meta.items() if k.isdigit()}
# elif isinstance(meta, list):
#     print("📂 Detected format: list")
#     id2skill = {int(s["id"]): s for s in meta if isinstance(s, dict) and "id" in s}
# else:
#     raise ValueError("❌ Unsupported meta format")

# print(f"✅ Loaded {len(id2skill)} ESCO skills into memory")

# # === Models ===
# model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
# try:
#     nlp = spacy.load("en_core_web_sm")  # lightweight for tokenization
# except OSError:
#     raise RuntimeError("⚠️ Run: python -m spacy download en_core_web_sm")


# # === Helpers ===
# def extract_text_from_pdf(pdf_path):
#     reader = PdfReader(pdf_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text


# def tokenize_text(text):
#     """Use spaCy for tokenization + multiword noun chunks."""
#     doc = nlp(text)
#     tokens = []

#     # single words
#     for token in doc:
#         t = token.text.strip().lower()
#         t = re.sub(r"[\n\r,;]+", " ", t).strip()
#         if not t or len(t) < 2:
#             continue
#         if len(t.split()) > 3:  # skip long junk
#             continue
#         if not re.match(r"^[a-z0-9\-\+#\.]+$", t):
#             continue
#         tokens.append(t)

#     # multiword noun chunks
#     for chunk in doc.noun_chunks:
#         c = chunk.text.strip().lower()
#         c = re.sub(r"[\n\r,;]+", " ", c).strip()
#         if 2 <= len(c.split()) <= 4:
#             tokens.append(c)

#     return list(set(tokens))


# def classify_skill(skill_data: dict) -> str:
#     """
#     Dynamically classify based on ESCO metadata instead of hardcoded rules.
#     Looks for 'skillType', 'conceptType', 'broaderConcept'.
#     """
#     for key in ("skillType", "conceptType", "broaderConcept"):
#         if key in skill_data and isinstance(skill_data[key], str):
#             val = skill_data[key].lower()
#             if any(x in val for x in ["ict", "digital", "technical", "engineering", "programming", "knowledge"]):
#                 return "Technical Skills"
#             if any(x in val for x in ["management", "business", "entrepreneurship", "administration", "strategy", "marketing"]):
#                 return "Management Skills"
#             if any(x in val for x in ["communication", "personal", "soft", "attitude", "collaboration", "teamwork"]):
#                 return "Soft Skills"
#     return "Other"


# def normalize_name(name: str) -> str:
#     """Normalize casing for nicer JSON output (title case except acronyms)."""
#     if not name:
#         return ""
#     words = []
#     for w in name.split():
#         if w.isupper():
#             words.append(w)  # keep acronyms like SQL, CSS
#         elif w.lower() in ["sql", "css", "html", "javascript", "python", "java"]:
#             words.append(w.upper() if w != "Python" else "Python")
#         else:
#             words.append(w.capitalize())
#     return " ".join(words)


# def normalize_skills(tokens):
#     results = {}
#     embeddings = model.encode(tokens, convert_to_numpy=True, normalize_embeddings=True)

#     D, I = index.search(embeddings, 3)  # top-3 ESCO matches
#     for token, scores, ids in zip(tokens, D, I):
#         best_id = int(ids[0])
#         best_score = float(scores[0])
#         skill_data = id2skill.get(best_id, {})

#         if debug:
#             print(f"\n🔎 Token: {token}")
#             for idx, sc in zip(ids, scores):
#                 sd = id2skill.get(int(idx), {})
#                 print(f"   -> {sc:.3f} | {sd.get('preferredLabel')} ({sd.get('uri')})")

#         if best_score >= threshold and skill_data:
#             esco_skill = normalize_name(skill_data.get("preferredLabel", ""))
#             uri = skill_data.get("uri", "")
#             category = classify_skill(skill_data)

#             # Deduplicate: keep highest score for same skill
#             key = (esco_skill, category)
#             if key not in results or best_score > results[key]["score"]:
#                 results[key] = {
#                     "skill": esco_skill,
#                     "score": round(best_score, 3),
#                     "uri": uri,
#                     "category": category,
#                 }

#     return list(results.values())


# def group_by_category(skills):
#     grouped = defaultdict(list)
#     for s in skills:
#         grouped[s["category"]].append({
#             "skill": s["skill"],
#             "score": s["score"],
#             "uri": s["uri"]
#         })
#     return grouped


# def main(pdf_path):
#     text = extract_text_from_pdf(pdf_path)
#     tokens = tokenize_text(text)

#     print(f"\n📑 Extracted {len(tokens)} clean tokens from resume")

#     skills = normalize_skills(tokens)
#     grouped_skills = group_by_category(skills)

#     output = {
#         "resume_name": os.path.basename(pdf_path),
#         "skills_by_category": grouped_skills,
#     }

#     out_path = os.path.splitext(pdf_path)[0] + "_skills.json"
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(output, f, indent=2)

#     print(f"\n✅ Results saved to {out_path}")


# if __name__ == "__main__":
#     import sys
#     pdf_path = sys.argv[1]
#     main(pdf_path)


# parse_resume.py
# import os
# import sys
# import json
# import faiss
# import numpy as np
# import spacy
# from PyPDF2 import PdfReader
# from sentence_transformers import SentenceTransformer


# # === Paths ===
# index_path = "vector_store/esco_faiss.index"
# meta_path = "vector_store/esco_meta.json"

# # === Load spaCy model ===
# try:
#     nlp = spacy.load("en_core_web_sm")
# except OSError:
#     print("⚠️ spaCy model not found. Run: python -m spacy download en_core_web_sm")
#     sys.exit(1)

# # === Load FAISS index + metadata ===
# print("📥 Loading FAISS index + metadata...")
# index = faiss.read_index(index_path)

# with open(meta_path, "r", encoding="utf-8") as f:
#     meta = json.load(f)

# id2skill = {int(k): v for k, v in meta.items()}

# # === Embedding model ===
# model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


# # === Resume text extraction ===
# def extract_text_from_pdf(path):
#     reader = PdfReader(path)
#     text = []
#     for page in reader.pages:
#         t = page.extract_text()
#         if t:
#             text.append(t)
#     return "\n".join(text)


# # === Candidate token extraction ===
# def extract_candidate_phrases(text):
#     doc = nlp(text)
#     phrases = set()

#     # Individual tokens
#     for token in doc:
#         if token.is_stop or token.is_punct or len(token.text) < 2:
#             continue
#         phrases.add(token.lemma_.strip().lower())

#     # Multi-word noun chunks
#     for chunk in doc.noun_chunks:
#         chunk_text = chunk.text.strip().lower()
#         if len(chunk_text.split()) > 1:
#             phrases.add(chunk_text)

#     return list(phrases)


# # === Phrase deduplication (remove overlaps) ===
# def deduplicate_phrases(phrases):
#     """
#     Removes overlapping phrases: 
#     If 'machine' and 'machine learning' both appear, keep only 'machine learning'.
#     """
#     # Sort so longer, multi-word phrases come first
#     phrases = sorted(set(phrases), key=lambda x: (-len(x.split()), -len(x)))
#     filtered = []

#     for phrase in phrases:
#         if not any(phrase in longer for longer in filtered if phrase != longer):
#             filtered.append(phrase)

#     return filtered


# # === Match tokens to ESCO ===
# def match_to_esco(phrases, top_k=1, threshold=0.6):
#     skills = []

#     for phrase in phrases:
#         emb = model.encode(phrase, normalize_embeddings=True).astype("float32")
#         emb = np.expand_dims(emb, axis=0)

#         scores, idxs = index.search(emb, top_k)

#         for score, idx in zip(scores[0], idxs[0]):
#             if score < threshold:
#                 continue
#             if idx not in id2skill:
#                 continue

#             esco_entry = id2skill[idx]
#             cats = esco_entry.get("categories", {})

#             skills.append({
#                 "skill": esco_entry["preferredLabel"].strip().title(),
#                 "score": round(float(score), 3),
#                 "uri": esco_entry.get("uri", ""),
#                 "subcategories": [
#                     cats.get("level_0", ""),
#                     cats.get("level_1", ""),
#                     cats.get("level_2", ""),
#                     cats.get("level_3", "")
#                 ]
#             })

#     return skills


# # === Group by Level 0 ===
# def group_by_category(skills):
#     grouped = {}
#     for s in skills:
#         level_0 = s.get("subcategories", [])[0] or "Other"

#         grouped.setdefault(level_0, []).append({
#             "skill": s["skill"],
#             "score": s["score"],
#             "uri": s["uri"],
#             "subcategories": [c for c in s["subcategories"] if c]
#         })
#     return grouped


# # === Main ===
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python parse_resume.py <resume.pdf>")
#         sys.exit(1)

#     resume_path = sys.argv[1]
#     resume_name = os.path.basename(resume_path)

#     print(f"📂 Parsing resume: {resume_name}")

#     text = extract_text_from_pdf(resume_path)
#     phrases = extract_candidate_phrases(text)

#     print(f"📑 Extracted {len(phrases)} candidate phrases")

#     # 🔥 Apply deduplication
#     phrases = deduplicate_phrases(phrases)
#     print(f"✨ After deduplication: {len(phrases)} phrases remain")

#     matched = match_to_esco(phrases, top_k=1, threshold=0.6)

#     # Deduplicate matches (keep highest score per skill+subcategory combo)
#     dedup = {}
#     for m in matched:
#         key = (m["skill"], tuple(m.get("subcategories", [])))
#         if key not in dedup or m["score"] > dedup[key]["score"]:
#             dedup[key] = m

#     grouped = group_by_category(list(dedup.values()))

#     output = {
#         "resume_name": resume_name,
#         "skills_by_category": grouped
#     }

#     out_path = os.path.splitext(resume_path)[0] + "_skills.json"
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(output, f, indent=2, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# import os
# import json
# import faiss
# import fitz  # PyMuPDF for PDF text extraction
# import numpy as np
# from sentence_transformers import SentenceTransformer

# # -------------------------------
# # Utility Functions
# # -------------------------------

# def extract_text_from_pdf(pdf_path):
#     """Extract raw text from PDF resume using PyMuPDF."""
#     text = ""
#     with fitz.open(pdf_path) as doc:
#         for page in doc:
#             text += page.get_text("text") + "\n"
#     return text


# def detect_candidate_skills(text, keywords=None):
#     """
#     Very simple skill detector.
#     Later this can be replaced by spaCy NER or regex patterns.
#     """
#     if keywords is None:
#         keywords = ["Python", "SQL", "Marketing", "Strategic", "Management", "Research"]

#     detected = []
#     for kw in keywords:
#         if kw.lower() in text.lower():
#             detected.append(kw)
#     return detected


# def load_esco_resources():
#     """Load FAISS index and metadata."""
#     script_dir = os.path.dirname(os.path.abspath(__file__))
#     project_root = os.path.dirname(script_dir)

#     index_path = os.path.join(project_root, "vector_store", "esco_faiss.index")
#     meta_path = os.path.join(project_root, "vector_store", "esco_meta.json")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)

#     return index, meta


# def map_to_esco(skill_text, model, index, meta, top_k=1):
#     """Embed skill text with CareerBERT and map to closest ESCO skill."""
#     q_vec = model.encode([skill_text], normalize_embeddings=True).astype("float32")
#     scores, idxs = index.search(q_vec, top_k)

#     results = []
#     for s, i in zip(scores[0], idxs[0]):
#         result = meta[str(i)].copy()
#         result["score"] = float(s)
#         results.append(result)
#     return results


# # -------------------------------
# # Main Resume Parsing Function
# # -------------------------------

# def parse_resume(pdf_path, keywords=None):
#     """Parse resume, detect skills, map to ESCO using CareerBERT + FAISS."""
#     print(f"📄 Parsing resume: {pdf_path}")

#     # Load models and resources
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")
#     index, meta = load_esco_resources()

#     # Extract raw text
#     raw_text = extract_text_from_pdf(pdf_path)

#     # Detect candidate skills
#     detected_skills = detect_candidate_skills(raw_text, keywords=keywords)
#     print(f"✅ Detected {len(detected_skills)} candidate skills.")

#     # Map detected skills to ESCO
#     categorized = {}
#     for skill in detected_skills:
#         esco_matches = map_to_esco(skill, model, index, meta, top_k=1)
#         if not esco_matches:
#             continue

#         best_match = esco_matches[0]
#         cats = list(best_match.get("categories", {}).values())
#         category_name = cats[0] if cats else "uncategorized"

#         if category_name not in categorized:
#             categorized[category_name] = []

#         categorized[category_name].append({
#             "skill": best_match["preferredLabel"],
#             "score": round(best_match["score"], 3),
#             "uri": best_match["uri"],
#             "subcategories": cats
#         })

#     # Build final JSON
#     output = {
#         "resume_name": os.path.basename(pdf_path),
#         "skills_by_category": categorized
#     }
#     return output


# # -------------------------------
# # Run as Script
# # -------------------------------
# if __name__ == "__main__":
#     script_dir = os.path.dirname(os.path.abspath(__file__))
#     sample_resume = os.path.join(script_dir, "..", "data", "resumes", "sample.pdf")

#     results = parse_resume(sample_resume)
#     out_path = os.path.join(script_dir, "..", "data", "resumes", "sample_skills.json")

#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=2, ensure_ascii=False)

#     print(f"\n🎉 Resume parsing complete. Output saved to {out_path}")



# import os
# import sys
# import json
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# import faiss
# import numpy as np
# from PyPDF2 import PdfReader

# class SkillGapAnalyzer:
#     """
#     Analyzes a resume to identify top skills and find matching ESCO occupations
#     based on skill overlap. Uses a consistent CareerBERT embedding model.
#     """
#     def __init__(self, project_root):
#         self.project_root = project_root
#         self.data_dir = os.path.join(project_root, "data")
#         self.vector_store_dir = os.path.join(project_root, "vector_store")
        
#         # --- 1. Load the CareerBERT Model ---
#         model_name = "lwolfrum2/careerbert-jg"
#         print(f"📥 Loading SentenceTransformer model: {model_name}...")
#         self.model = SentenceTransformer(model_name)
        
#         # --- 2. Load the FAISS Index and Metadata ---
#         self._load_esco_skill_index()
        
#         # --- 3. Load Occupation Data ---
#         self._load_occupations()

#     def _load_esco_skill_index(self):
#         """Loads the FAISS index and the skill metadata JSON."""
#         print("📥 Loading FAISS index and skill metadata...")
#         index_path = os.path.join(self.vector_store_dir, "esco_faiss.index")
#         meta_path = os.path.join(self.vector_store_dir, "esco_meta.json")
#         try:
#             self.index = faiss.read_index(index_path)
#             with open(meta_path, "r", encoding="utf-8") as f:
#                 self.skill_meta = json.load(f)
#         except FileNotFoundError:
#             print(f"❌ ERROR: Missing FAISS index or metadata in '{self.vector_store_dir}'.")
#             print("Please run 'build_esco_index.py' first.")
#             sys.exit(1)

#     def _load_occupations(self):
#         """Loads all occupation-related CSVs and builds the necessary mappings."""
#         print("📥 Loading occupation datasets...")
#         # Corrected path to look inside the 'data/esco' subfolder
#         esco_data_dir = os.path.join(self.data_dir, "esco")
#         try:
#             occ_df = pd.read_csv(os.path.join(esco_data_dir, "occupations_en.csv"), dtype=str).fillna("")
#             occ_skill_df = pd.read_csv(os.path.join(esco_data_dir, "occupationSkillRelations_en.csv"), dtype=str).fillna("")
#             isco_df = pd.read_csv(os.path.join(esco_data_dir, "ISCOGroups_en.csv"), dtype=str).fillna("")
#         except FileNotFoundError as e:
#             print(f"❌ ERROR: Could not find occupation data file: {e.filename}")
#             print("Please ensure your ESCO CSV files are in the 'data/esco' folder.")
#             sys.exit(1)
            
#         # Build occupation URI -> required skill URIs map
#         self.occ_to_skills = {}
#         for _, row in occ_skill_df.iterrows():
#             occ_uri = row.get("occupationUri", "").strip()
#             skill_uri = row.get("skillUri", "").strip()
#             if occ_uri and skill_uri:
#                 self.occ_to_skills.setdefault(occ_uri, set()).add(skill_uri)

#         # Build ISCO group URI -> label map
#         isco_map = {row["conceptUri"]: row.get("preferredLabel", "") for _, row in isco_df.iterrows() if row.get("conceptUri")}

#         # Build occupation URI -> metadata map
#         self.occ_meta = {}
#         for _, row in occ_df.iterrows():
#             occ_uri = row.get("conceptUri", "").strip()
#             if occ_uri:
#                 self.occ_meta[occ_uri] = {
#                     "title": row.get("preferredLabel", ""),
#                     "description": row.get("description", ""),
#                     "iscoGroup": isco_map.get(row.get("iscoGroup", ""), ""),
#                 }

#     def extract_text_from_resume(self, resume_path: str) -> str:
#         """Extracts text content from a PDF or TXT file."""
#         print(f"📄 Reading text from '{os.path.basename(resume_path)}'...")
#         if resume_path.lower().endswith(".pdf"):
#             reader = PdfReader(resume_path)
#             return " ".join([page.extract_text() or "" for page in reader.pages])
#         else:
#             with open(resume_path, "r", encoding="utf-8", errors="ignore") as f:
#                 return f.read()

#     def analyze(self, resume_path: str):
#         """
#         Performs the full analysis pipeline on a given resume file.
#         This version uses a segmentation strategy for more accurate skill extraction.
#         """
#         # 1. Extract text and create meaningful segments
#         resume_text = self.extract_text_from_resume(resume_path)
#         segments = [seg.strip() for seg in resume_text.split('\n') if len(seg.strip().split()) > 3]
#         print(f"📄 Segmented resume into {len(segments)} meaningful chunks.")

#         if not segments:
#             print("⚠️  Could not find enough text segments to analyze.")
#             return {"top_10_skills_identified": [], "top_5_matching_occupations": []}

#         # 2. Encode all segments in a single batch
#         print("🤖 Finding top skills for each resume segment...")
#         query_vectors = self.model.encode(segments, normalize_embeddings=True, show_progress_bar=True).astype("float32")
        
#         # 3. Search the FAISS index for the top 5 skills for each segment
#         scores, indices = self.index.search(query_vectors, k=5)
        
#         # 4. Aggregate, deduplicate, and rank the results
#         unique_skills = {}
#         for i, segment_indices in enumerate(indices):
#             for j, idx in enumerate(segment_indices):
#                 score = scores[i][j]
#                 skill_info = self.skill_meta.get(str(idx))
#                 if skill_info:
#                     uri = skill_info["uri"]
#                     # If skill is new or this is a better score, add/update it
#                     if uri not in unique_skills or score > unique_skills[uri]["score"]:
#                         unique_skills[uri] = {"score": score, "info": skill_info}
        
#         # Sort all unique skills by their best score
#         sorted_skills = sorted(unique_skills.values(), key=lambda x: x["score"], reverse=True)
#         top_10_skills = [s["info"] for s in sorted_skills[:10]]
        
#         # 5. Find matching occupations based on the new, more accurate skill list
#         print("🤝 Calculating occupation overlap scores...")
#         skill_uris_from_resume = {s["uri"] for s in top_10_skills}
        
#         occ_scores = []
#         for occ_uri, required_skills in self.occ_to_skills.items():
#             overlap = skill_uris_from_resume.intersection(required_skills)
#             if overlap:
#                 occ_scores.append((occ_uri, len(overlap)))
        
#         # 6. Get the top 5 occupations with the highest overlap
#         top_5_occupations_sorted = sorted(occ_scores, key=lambda x: x[1], reverse=True)[:5]
        
#         top_occupations = []
#         for occ_uri, score in top_5_occupations_sorted:
#             meta = self.occ_meta.get(occ_uri, {})
#             top_occupations.append({
#                 "title": meta.get("title", "N/A"),
#                 "description": meta.get("description", "N/A"),
#                 "iscoGroup": meta.get("iscoGroup", "N/A"),
#                 "matching_skills_count": score
#             })
            
#         return {
#             "top_10_skills_identified": top_10_skills,
#             "top_5_matching_occupations": top_occupations
#         }


# def main():
#     if len(sys.argv) < 2:
#         print("Usage: python app/skill_gap_analyzer.py <path_to_resume>")
#         sys.exit(1)
        
#     resume_path = sys.argv[1]
    
#     # Dynamically determine the project root from the script's location
#     script_dir = os.path.dirname(os.path.abspath(__file__))
#     project_root = os.path.dirname(script_dir)
    
#     # Initialize the analyzer (this will load all models and data)
#     analyzer = SkillGapAnalyzer(project_root)
    
#     # Run the analysis
#     results = analyzer.analyze(resume_path)
    
#     # Define and create the output path
#     output_filename = os.path.splitext(os.path.basename(resume_path))[0] + "_skills.json"
#     output_path = os.path.join(project_root, "data", "resumes", output_filename)
#     os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
#     # Save the results to the specified JSON file
#     with open(output_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)
        
#     print(f"\n🎉 Success! Analysis complete. Results saved to '{output_path}'")

# if __name__ == "__main__":
#     main()






# # app/parse_resume.py

# import os
# import sys
# import json
# import faiss
# import numpy as np
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader


# # ---------- Load Resume Text ----------
# def load_resume_text(pdf_path: str) -> str:
#     reader = PdfReader(pdf_path)
#     text = " ".join([page.extract_text() or "" for page in reader.pages])
#     return text.strip()


# # ---------- Load ESCO Skill Index ----------
# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")

#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)
#     return index, meta


# # ---------- Load Occupations + ISCO + Relations ----------
# def load_occupations(project_root: str):
#     occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
#     isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
#     relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

#     occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
#     isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
#     rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

#     # Merge ISCO group info into occupations
#     occ_df = occ_df.merge(
#         isco_df[["code", "preferredLabel", "description"]],
#         left_on="iscoGroup",
#         right_on="code",
#         how="left",
#         suffixes=("", "_isco"),
#     )

#     # Build skillUris list per occupation
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

#     return occ_df


# # ---------- Match Skills from Resume ----------
# def extract_top_skills(resume_text: str, model, index, meta, top_k=10):
#     emb = model.encode([resume_text], normalize_embeddings=True)
#     D, I = index.search(emb, top_k)
#     skills = []
#     for idx in I[0]:
#         if str(idx) in meta:
#             skills.append(meta[str(idx)])
#     return skills


# # ---------- Match Occupations ----------
# def match_occupations(skills, occ_df, top_k=5):
#     skill_uris = {s["uri"] for s in skills}
#     occ_matches = []

#     for _, row in occ_df.iterrows():
#         occ_skill_uris = set(row["skillUris"])
#         overlap = skill_uris.intersection(occ_skill_uris)
#         if overlap:
#             occ_matches.append({
#                 "title": row["preferredLabel"],
#                 "description": row["description"],
#                 "iscoGroup": f"{row['code']} - {row['preferredLabel_isco']}" if row["code"] else "",
#                 "matching_skills_count": len(overlap)
#             })

#     occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
#     return occ_matches[:top_k]


# # ---------- Main Pipeline ----------
# def parse_resume(pdf_path: str, project_root: str):
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume...")
#     resume_text = load_resume_text(pdf_path)

#     print("📦 Loading ESCO index...")
#     index, meta = load_esco_index(project_root)

#     print("👔 Loading occupations + ISCO groups...")
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting top skills...")
#     top_skills = extract_top_skills(resume_text, model, index, meta, top_k=10)

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, top_k=5)

#     results = {
#         "top_10_skills_identified": top_skills,
#         "top_5_matching_occupations": top_occupations
#     }

#     out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# # ---------- Run ----------
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <resume.pdf>")
#         sys.exit(1)

#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     pdf_path = sys.argv[1]
#     parse_resume(pdf_path, project_root)


# import os
# import sys
# import json
# import faiss
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader


# # ---------- Load Resume Text ----------
# def load_resume_text(pdf_path: str) -> str:
#     reader = PdfReader(pdf_path)
#     text = " ".join([page.extract_text() or "" for page in reader.pages])
#     return text.strip()


# # ---------- Load ESCO Skill Index ----------
# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")

#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)

#     # Reverse lookup: conceptUri → metadata
#     uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
#     return index, meta, uri_to_meta


# # ---------- Load Occupations + ISCO + Relations ----------
# def load_occupations(project_root: str):
#     occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
#     isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
#     relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

#     occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
#     isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
#     rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

#     # Merge ISCO group info into occupations
#     occ_df = occ_df.merge(
#         isco_df[["code", "preferredLabel", "description"]],
#         left_on="iscoGroup",
#         right_on="code",
#         how="left",
#         suffixes=("", "_isco"),
#     )

#     # Build skillUris list per occupation
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

#     return occ_df


# # ---------- Match Skills from Resume ----------
# def extract_top_skills(resume_text: str, model, index, meta, top_k=10):
#     emb = model.encode([resume_text], normalize_embeddings=True)
#     D, I = index.search(emb, top_k)
#     skills = []
#     for idx in I[0]:
#         if str(idx) in meta:
#             entry = meta[str(idx)]
#             skills.append({
#                 "preferredLabel": entry.get("preferredLabel", ""),
#                 "conceptUri": entry.get("uri", ""),
#                 "description": entry.get("description", "")
#             })
#     return skills


# # ---------- Match Occupations ----------
# def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
#     skill_uris = {s["conceptUri"] for s in skills}
#     occ_matches = []

#     for _, row in occ_df.iterrows():
#         occ_skill_uris = set(row["skillUris"])
#         overlap = skill_uris.intersection(occ_skill_uris)
#         if overlap:
#             # Required skills with labels from meta
#             required_skills = [
#                 {"label": uri_to_meta[uri]["preferredLabel"], "conceptUri": uri}
#                 for uri in occ_skill_uris if uri in uri_to_meta
#             ]
#             matched_skills = [s for s in required_skills if s["conceptUri"] in overlap]
#             missing_skills = [s for s in required_skills if s["conceptUri"] not in skill_uris]

#             occ_matches.append({
#                 "preferredLabel": row["preferredLabel"],
#                 "conceptUri": row["conceptUri"],
#                 "iscoGroup": {
#                     "code": row["code"],
#                     "name": row["preferredLabel_isco"],
#                     "description": row["description"]
#                 },
#                 "matching_skills_count": len(overlap),
#                 "required_skills": required_skills,
#                 "matched_skills": matched_skills,
#                 "skill_gap": missing_skills
#             })

#     occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
#     return occ_matches[:top_k]


# # ---------- Main Pipeline ----------
# def parse_resume(pdf_path: str, project_root: str):
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume...")
#     resume_text = load_resume_text(pdf_path)

#     print("📦 Loading ESCO index...")
#     index, meta, uri_to_meta = load_esco_index(project_root)

#     print("👔 Loading occupations + ISCO groups...")
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting top skills...")
#     top_skills = extract_top_skills(resume_text, model, index, meta, top_k=10)

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, uri_to_meta, top_k=5)

#     results = {
#         "top_10_skills_identified": top_skills,
#         "top_5_matching_occupations": top_occupations
#     }

#     out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# # ---------- Run ----------
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <resume.pdf>")
#         sys.exit(1)

#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     pdf_path = sys.argv[1]
#     parse_resume(pdf_path, project_root)

# import os
# import sys
# import json
# import faiss
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader


# # ---------- Load Resume Text ----------
# def load_resume_text(pdf_path: str) -> str:
#     reader = PdfReader(pdf_path)
#     text = " ".join([page.extract_text() or "" for page in reader.pages])
#     return text.strip()


# # ---------- Load ESCO Skill Index ----------
# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")

#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)

#     # Reverse lookup: conceptUri → metadata
#     uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
#     return index, meta, uri_to_meta


# # ---------- Load Occupations + ISCO + Relations ----------
# def load_occupations(project_root: str):
#     occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
#     isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
#     relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

#     occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
#     isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
#     rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

#     # Merge ISCO group info into occupations
#     occ_df = occ_df.merge(
#         isco_df[["code", "preferredLabel", "description"]],
#         left_on="iscoGroup",
#         right_on="code",
#         how="left",
#         suffixes=("", "_isco"),
#     )

#     # Build skillUris list per occupation
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

#     return occ_df


# # ---------- Match Skills from Resume ----------
# def extract_top_skills(resume_text: str, model, index, meta, top_k=10):
#     emb = model.encode([resume_text], normalize_embeddings=True)
#     D, I = index.search(emb, top_k)
#     skills = []
#     for idx in I[0]:
#         if str(idx) in meta:
#             entry = meta[str(idx)]
#             skills.append({
#                 "preferredLabel": entry.get("preferredLabel", ""),
#                 "conceptUri": entry.get("uri", ""),
#                 "description": entry.get("description", "")
#             })
#     return skills


# # ---------- Match Occupations ----------
# def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
#     skill_uris = {s["conceptUri"] for s in skills}
#     occ_matches = []

#     for _, row in occ_df.iterrows():
#         occ_skill_uris = set(row["skillUris"])
#         overlap = skill_uris.intersection(occ_skill_uris)
#         if overlap:
#             matched_labels = [uri_to_meta[uri]["preferredLabel"] for uri in overlap if uri in uri_to_meta]
#             occ_matches.append({
#                 "preferredLabel": row["preferredLabel"],
#                 "conceptUri": row["conceptUri"],
#                 "iscoGroup": {
#                     "code": row["code"],
#                     "name": row["preferredLabel_isco"],
#                     "description": row["description"]
#                 },
#                 "matching_skills_count": len(overlap),
#                 "matched_skills": matched_labels,
#                 "required_skills_count": len(occ_skill_uris)
#             })

#     occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
#     return occ_matches[:top_k]


# # ---------- Step 5: Skill Gap Analysis ----------
# def skill_gap_summary(top_skills, chosen_role, uri_to_meta):
#     student_uris = {s["conceptUri"]: s["preferredLabel"] for s in top_skills}
#     required_uris = set(chosen_role.get("skillUris", []))

#     overlap = set(student_uris.keys()).intersection(required_uris)
#     gaps = required_uris - set(student_uris.keys())

#     summary = {
#         "strengths": [student_uris[u] for u in overlap if u in student_uris],
#         "gaps": [uri_to_meta[u]["preferredLabel"] for u in gaps if u in uri_to_meta],
#         "underdeveloped": []  # placeholder
#     }
#     return summary


# # ---------- Main Pipeline ----------
# def parse_resume(pdf_path: str, project_root: str):
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume...")
#     resume_text = load_resume_text(pdf_path)

#     print("📦 Loading ESCO index...")
#     index, meta, uri_to_meta = load_esco_index(project_root)

#     print("👔 Loading occupations + ISCO groups...")
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting top skills...")
#     top_skills = extract_top_skills(resume_text, model, index, meta, top_k=10)

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, uri_to_meta, top_k=5)

#     # Step 5: Gap analysis on best role (if available)
#     chosen_role_uri = top_occupations[0]["conceptUri"] if top_occupations else None
#     chosen_role_row = occ_df.loc[occ_df["conceptUri"] == chosen_role_uri].iloc[0] if chosen_role_uri else None
#     skill_gap = skill_gap_summary(top_skills, chosen_role_row if chosen_role_row is not None else {}, uri_to_meta)

#     results = {
#         "top_10_skills_identified": top_skills,
#         "top_5_matching_occupations": top_occupations,
#         "skill_gap_analysis": skill_gap
#     }

#     out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# # ---------- Run ----------
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <resume.pdf>")
#         sys.exit(1)

#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     pdf_path = sys.argv[1]
#     parse_resume(pdf_path, project_root)



# import os
# import sys
# import json
# import re
# import faiss
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader

# # All helper functions (load_resume_text, load_esco_index, etc.) are included here

# def load_resume_text(pdf_path: str) -> str:
#     try:
#         reader = PdfReader(pdf_path)
#         text = "\n".join([page.extract_text() or "" for page in reader.pages])
#         return text.strip()
#     except Exception as e:
#         print(f"Error reading PDF {pdf_path}: {e}")
#         sys.exit(1)

# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")
#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("Missing ESCO FAISS index or metadata. Run app/build_esco_index.py first.")
#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)
#     uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
#     return index, meta, uri_to_meta

# def load_occupations(project_root: str):
#     data_dir = os.path.join(project_root, "data", "esco")
#     occ_df = pd.read_csv(os.path.join(data_dir, "occupations_en.csv"), dtype=str).fillna("")
#     isco_df = pd.read_csv(os.path.join(data_dir, "ISCOGroups_en.csv"), dtype=str).fillna("")
#     rel_df = pd.read_csv(os.path.join(data_dir, "occupationSkillRelations_en.csv"), dtype=str).fillna("")
#     occ_df = occ_df.merge(isco_df[["code", "preferredLabel", "description"]], left_on="iscoGroup", right_on="code", how="left", suffixes=("", "_isco"))
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])
#     return occ_df

# def extract_top_skills(resume_text: str, model, index, meta, similarity_threshold=0.75):
#     skills_block = ""
#     try:
#         start_match = re.search(r'S\s*K\s*I\s*L\s*L\s*S', resume_text, re.IGNORECASE)
#         if not start_match: return []
#         start_index = start_match.end()
#         end_match = re.search(r'P\s*O\s*S\s*I\s*T\s*I\s*O\s*N\s*S\s*O\s*F\s*R\s*E\s*S\s*P\s*O\s*N\s*S\s*I\s*B\s*I\s*L\s*I\s*T\s*Y|C\s*E\s*R\s*T\s*I\s*F\s*I\s*C\s*A\s*T\s*I\s*O\s*N|A\s*C\s*H\s*I\s*E\s*V\s*E\s*M\s*E\s*N\s*T', resume_text[start_index:], re.IGNORECASE)
#         skills_block = resume_text[start_index : start_index + end_match.start()] if end_match else resume_text[start_index : start_index + 500]
#     except: return []

#     clean_text = re.sub(r'(technical|additional)\s*skills\s*:?', '', skills_block, flags=re.IGNORECASE)
#     clean_text = clean_text.replace('Power BI', 'PowerBI').replace('Scikit-learn', 'Scikitlearn')
#     candidate_skills = [skill.strip() for skill in re.split(r',|\n|•|-|/', clean_text) if skill.strip() and len(skill.strip()) > 1]
#     stop_words = {'and', 'with', 'various', 'using', 'libraries', 'analysis', 'thinking', 'data'}
#     candidate_skills = [s for s in candidate_skills if s.lower() not in stop_words and len(s.split()) < 4]
    
#     if not candidate_skills: return []
#     print(f"ℹ️  Found {len(candidate_skills)} candidate skills to analyze.")

#     embeddings = model.encode(candidate_skills, normalize_embeddings=True)
    
#     # --- THIS IS THE CRITICAL FIX ---
#     # For IndexFlatIP, FAISS returns the dot product scores directly.
#     # We pass k=len(embeddings) to get all scores, then process them.
#     D, I = index.search(embeddings, 1) 
    
#     matched_skills = {}
#     for i, candidate in enumerate(candidate_skills):
#         idx = I[i][0]
#         # For an IP index, the score 'D' is the similarity score itself.
#         # No '1 - D' calculation is needed.
#         similarity = D[i][0] 

#         if similarity >= similarity_threshold:
#             entry = meta.get(str(idx), {})
#             uri = entry.get("uri")
#             if uri and uri not in matched_skills:
#                 matched_skills[uri] = {"preferredLabel": entry.get("preferredLabel", ""), "conceptUri": uri, "description": entry.get("description", ""), "matched_from": candidate, "similarity": round(float(similarity), 4)}
#     return sorted(matched_skills.values(), key=lambda x: x["similarity"], reverse=True)

# def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
#     skill_uri_to_score = {s["conceptUri"]: s["similarity"] for s in skills}
#     skill_uris = set(skill_uri_to_score.keys())
#     occ_matches = []
#     for _, row in occ_df.iterrows():
#         overlap = skill_uris.intersection(set(row["skillUris"]))
#         if overlap:
#             match_score = sum(skill_uri_to_score[uri] for uri in overlap)
#             occ_matches.append({"preferredLabel": row["preferredLabel"], "conceptUri": row["conceptUri"], "iscoGroup": {"code": row.get("code", ""), "name": row.get("preferredLabel_isco", ""), "description": row.get("description_isco", "")}, "matching_skills_count": len(overlap), "match_score": round(match_score, 4), "matched_skills": [uri_to_meta[uri]["preferredLabel"] for uri in overlap if uri in uri_to_meta], "required_skills_count": len(set(row["skillUris"]))})
#     return sorted(occ_matches, key=lambda x: x["match_score"], reverse=True)[:top_k]

# def skill_gap_summary(top_skills, chosen_role, uri_to_meta, resume_text):
#     # --- THIS IS THE CORRECTED LINE ---
#     if chosen_role is None or not isinstance(chosen_role, pd.Series):
#         return {"strengths": [], "gaps": [], "underdeveloped": []}
    
#     student_skills_map = {s["conceptUri"]: s for s in top_skills}
#     required_uris = set(chosen_role.get("skillUris", []))
#     overlap_uris = set(student_skills_map.keys()).intersection(required_uris)
#     gap_uris = required_uris - set(student_skills_map.keys())
    
#     strengths, underdeveloped = [], []
#     for uri in overlap_uris:
#         skill_info, skill_label = student_skills_map[uri], student_skills_map[uri]["preferredLabel"]
#         # A simple heuristic: if a skill's original keyword is in the resume more than once, 
#         # it suggests it was mentioned in the skills list and also in a project/experience.
#         if resume_text.lower().count(skill_info["matched_from"].lower()) > 1:
#             strengths.append(skill_label)
#         else:
#             underdeveloped.append(skill_label)
            
#     return {"strengths": sorted(strengths), "gaps": sorted([uri_to_meta[u]["preferredLabel"] for u in gap_uris if u in uri_to_meta]), "underdeveloped": sorted(underdeveloped)}
# def parse_resume(pdf_path: str, project_root: str):
#     print("⚙️  Loading the lwolfrum2/careerbert-jg model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume text...")
#     resume_text = load_resume_text(pdf_path)

#     print("📦 Loading ESCO index and data...")
#     index, meta, uri_to_meta = load_esco_index(project_root)
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting skills from resume...")
#     top_skills = extract_top_skills(resume_text, model, index, meta)

#     if not top_skills:
#         print("\n❌ No relevant skills identified. The similarity threshold may be too high.")
#     else:
#         print(f"✅ Found {len(top_skills)} relevant skills.")

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, uri_to_meta)

#     print("🔬 Performing skill gap analysis...")
#     chosen_role_row = None
#     if top_occupations:
#         chosen_role_df = occ_df.loc[occ_df["conceptUri"] == top_occupations[0]["conceptUri"]]
#         if not chosen_role_df.empty: chosen_role_row = chosen_role_df.iloc[0]
#     skill_gap = skill_gap_summary(top_skills, chosen_role_row, uri_to_meta, resume_text)

#     results = {"identified_skills": top_skills, "matching_occupations": top_occupations, "skill_gap_analysis": skill_gap}
    
#     output_filename = os.path.splitext(os.path.basename(pdf_path))[0] + "_skills.json"
#     out_path = os.path.join(project_root, "data", "resumes", output_filename)
#     os.makedirs(os.path.dirname(out_path), exist_ok=True)
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)
#     print(f"\n✅ Success! Results saved to {out_path}")

# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <path_to_your_resume.pdf>")
#         sys.exit(1)
#     pdf_path = sys.argv[1]
#     if not os.path.exists(pdf_path):
#         print(f"Error: The file '{pdf_path}' was not found.")
#         sys.exit(1)
#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     parse_resume(pdf_path, project_root)





# import os
# import sys
# import json
# import faiss
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader
# import spacy
# from spacy.matcher import PhraseMatcher
# from rapidfuzz import process  # for fuzzy matching


# # ---------- Load Resume Text ----------
# def load_resume_text(pdf_path: str) -> str:
#     reader = PdfReader(pdf_path)
#     text = " ".join([page.extract_text() or "" for page in reader.pages])
#     return text.strip()


# # ---------- Load ESCO Skill Index ----------
# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")

#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)

#     # Reverse lookup: conceptUri → metadata
#     uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
#     return index, meta, uri_to_meta


# # ---------- Load Occupations + ISCO + Relations ----------
# def load_occupations(project_root: str):
#     occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
#     isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
#     relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

#     occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
#     isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
#     rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

#     # Merge ISCO group info into occupations
#     occ_df = occ_df.merge(
#         isco_df[["code", "preferredLabel", "description"]],
#         left_on="iscoGroup",
#         right_on="code",
#         how="left",
#         suffixes=("", "_isco"),
#     )

#     # Build skillUris list per occupation
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

#     return occ_df


# # ---------- Detect Candidate Skills (NER + dictionary) ----------
# def detect_skills(resume_text: str, meta):
#     nlp = spacy.load("en_core_web_md")  # medium model for efficiency
#     matcher = PhraseMatcher(nlp.vocab, attr="LOWER")

#     # collect ESCO labels + synonyms
#     all_labels = []
#     for entry in meta.values():
#         all_labels.append(entry.get("preferredLabel", ""))
#         alt_labels = entry.get("altLabels", "")
#         if isinstance(alt_labels, str) and alt_labels:
#             all_labels.extend([x.strip() for x in alt_labels.split(",")])

#     patterns = [nlp.make_doc(label) for label in set(all_labels) if label]
#     matcher.add("SKILLS", patterns)

#     doc = nlp(resume_text)
#     matches = matcher(doc)

#     found = [doc[start:end].text for _, start, end in matches]
#     return list(set(found))


# # ---------- Normalize Mentions with Synonyms/Fuzzy Matching ----------
# def normalize_skill(skill, meta):
#     """
#     Given a detected skill mention, find the closest ESCO skill
#     using fuzzy string matching and return the full metadata entry.
#     """
#     labels_map = {v["preferredLabel"]: v for v in meta.values() if "preferredLabel" in v}
#     match, score, _ = process.extractOne(skill, labels_map.keys())
#     if match and score > 80:
#         return labels_map[match]  # full ESCO entry
#     return None


# # ---------- Extract Top Skills (Hybrid: NER + Fuzzy + Embedding Fallback) ----------
# import re

# def extract_top_skills(resume_text: str, model, index, meta, top_k=3):
#     sentences = re.split(r'(?<=[.!?])\s+', resume_text)
#     skills = []

#     for sent in sentences:
#         candidate_mentions = detect_skills(sent, meta)

#         if candidate_mentions:
#             # --- Normal pipeline ---
#             for mention in candidate_mentions:
#                 norm_entry = normalize_skill(mention, meta)
#                 query_label = norm_entry["preferredLabel"] if norm_entry else mention
#                 emb = model.encode([query_label], normalize_embeddings=True)
#                 D, I = index.search(emb, top_k)
#                 for idx in I[0]:
#                     if str(idx) in meta:
#                         skills.append(meta[str(idx)])
#         else:
#             # --- NEW: Embedding fallback for whole sentence ---
#             emb = model.encode([sent], normalize_embeddings=True)
#             D, I = index.search(emb, top_k)
#             for idx in I[0]:
#                 if str(idx) in meta:
#                     skills.append(meta[str(idx)])

#     # Deduplicate
#     seen = set()
#     unique_skills = []
#     for s in skills:
#         if s["uri"] not in seen:
#             unique_skills.append(s)
#             seen.add(s["uri"])

#     return unique_skills


# # ---------- Match Occupations ----------
# def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
#     skill_uris = {s["uri"] for s in skills}
#     occ_matches = []

#     for _, row in occ_df.iterrows():
#         occ_skill_uris = set(row["skillUris"])
#         overlap = skill_uris.intersection(occ_skill_uris)
#         if overlap:
#             matched_labels = [uri_to_meta[uri]["preferredLabel"] for uri in overlap if uri in uri_to_meta]
#             occ_matches.append({
#                 "preferredLabel": row["preferredLabel"],
#                 "conceptUri": row["conceptUri"],
#                 "iscoGroup": {
#                     "code": row["code"],
#                     "name": row["preferredLabel_isco"],
#                     "description": row["description"]
#                 },
#                 "matching_skills_count": len(overlap),
#                 "matched_skills": matched_labels,
#                 "required_skills_count": len(occ_skill_uris)
#             })

#     occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
#     return occ_matches[:top_k]


# # ---------- Step 5: Skill Gap Analysis ----------
# def skill_gap_summary(top_skills, chosen_role, uri_to_meta):
#     student_uris = {s["uri"]: s["preferredLabel"] for s in top_skills}
#     required_uris = set(chosen_role.get("skillUris", []))

#     overlap = set(student_uris.keys()).intersection(required_uris)
#     gaps = required_uris - set(student_uris.keys())

#     summary = {
#         "strengths": [student_uris[u] for u in overlap if u in student_uris],
#         "gaps": [uri_to_meta[u]["preferredLabel"] for u in gaps if u in uri_to_meta],
#         "underdeveloped": []  # placeholder
#     }
#     return summary


# # ---------- Main Pipeline ----------
# def parse_resume(pdf_path: str, project_root: str):
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume...")
#     resume_text = load_resume_text(pdf_path)

#     print("📦 Loading ESCO index...")
#     index, meta, uri_to_meta = load_esco_index(project_root)

#     print("👔 Loading occupations + ISCO groups...")
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting skills with hybrid method...")
#     top_skills = extract_top_skills(resume_text, model, index, meta, top_k=5)

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, uri_to_meta, top_k=5)

#     # Step 5: Gap analysis on best role (if available)
#     chosen_role_uri = top_occupations[0]["conceptUri"] if top_occupations else None
#     chosen_role_row = occ_df.loc[occ_df["conceptUri"] == chosen_role_uri].iloc[0] if chosen_role_uri else None
#     skill_gap = skill_gap_summary(top_skills, chosen_role_row if chosen_role_row is not None else {}, uri_to_meta)

#     results = {
#         "top_skills_identified": top_skills[:15],        # only first 15
#         "top_matching_occupations": top_occupations[:10], # only first 10
#         "skill_gap_analysis": skill_gap
#     }

#     out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# # ---------- Run ----------
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <resume.pdf>")
#         sys.exit(1)

#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     pdf_path = sys.argv[1]
#     parse_resume(pdf_path, project_root)

# import os
# import sys
# import json
# import faiss
# import pandas as pd
# from sentence_transformers import SentenceTransformer
# from PyPDF2 import PdfReader
# import spacy
# from spacy.matcher import PhraseMatcher
# from rapidfuzz import process  # for fuzzy matching
# import re
# import numpy as np
# from docx import Document   # for .docx resumes


# # ---------- Load Resume Text ----------
# def load_resume_text(file_path: str) -> str:
#     if file_path.lower().endswith(".pdf"):
#         reader = PdfReader(file_path)
#         text = " ".join([page.extract_text() or "" for page in reader.pages])
#     elif file_path.lower().endswith(".docx"):
#         doc = Document(file_path)
#         text = " ".join([para.text for para in doc.paragraphs])
#     else:
#         raise ValueError("Unsupported file format. Please use .pdf or .docx")
#     return text.strip()


# # ---------- Load ESCO Skill Index ----------
# def load_esco_index(project_root: str):
#     vector_store_dir = os.path.join(project_root, "vector_store")
#     index_path = os.path.join(vector_store_dir, "esco_faiss.index")
#     meta_path = os.path.join(vector_store_dir, "esco_meta.json")

#     if not os.path.exists(index_path) or not os.path.exists(meta_path):
#         raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

#     index = faiss.read_index(index_path)
#     with open(meta_path, "r", encoding="utf-8") as f:
#         meta = json.load(f)

#     # Reverse lookup: conceptUri → metadata
#     uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
#     return index, meta, uri_to_meta


# # ---------- Load Occupations + ISCO + Relations ----------
# def load_occupations(project_root: str):
#     occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
#     isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
#     relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

#     occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
#     isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
#     rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

#     # Merge ISCO group info into occupations
#     occ_df = occ_df.merge(
#         isco_df[["code", "preferredLabel", "description"]],
#         left_on="iscoGroup",
#         right_on="code",
#         how="left",
#         suffixes=("", "_isco"),
#     )

#     # Build skillUris list per occupation
#     occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
#     occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

#     return occ_df


# # ---------- Detect Candidate Skills (NER + dictionary) ----------
# def detect_skills(resume_text: str, meta):
#     nlp = spacy.load("en_core_web_md")  # medium model for efficiency
#     matcher = PhraseMatcher(nlp.vocab, attr="LOWER")

#     # collect ESCO labels + synonyms
#     all_labels = []
#     for entry in meta.values():
#         all_labels.append(entry.get("preferredLabel", ""))
#         alt_labels = entry.get("altLabels", "")
#         if isinstance(alt_labels, str) and alt_labels:
#             all_labels.extend([x.strip() for x in alt_labels.split(",")])

#     patterns = [nlp.make_doc(label) for label in set(all_labels) if label]
#     matcher.add("SKILLS", patterns)

#     doc = nlp(resume_text)
#     matches = matcher(doc)

#     found = [doc[start:end].text for _, start, end in matches]
#     return list(set(found))


# # ---------- Normalize Mentions with Synonyms/Fuzzy Matching ----------
# def normalize_skill(skill, meta):
#     labels_map = {v["preferredLabel"]: v for v in meta.values() if "preferredLabel" in v}
#     match, score, _ = process.extractOne(skill, labels_map.keys())
#     if match and score > 80:
#         return labels_map[match]  # full ESCO entry
#     return None


# # ---------- Extract Top Skills (Hybrid: NER + Fuzzy + Embedding Fallback) ----------
# def extract_top_skills(resume_text: str, model, index, meta, top_k=3):
#     sentences = re.split(r'(?<=[.!?])\s+', resume_text)
#     skills = []

#     for sent in sentences:
#         candidate_mentions = detect_skills(sent, meta)

#         if candidate_mentions:
#             for mention in candidate_mentions:
#                 norm_entry = normalize_skill(mention, meta)
#                 query_label = norm_entry["preferredLabel"] if norm_entry else mention
#                 emb = model.encode([query_label], normalize_embeddings=True)
#                 D, I = index.search(emb, top_k)
#                 for idx in I[0]:
#                     if str(idx) in meta:
#                         skills.append(meta[str(idx)])
#         else:
#             # fallback: semantic embedding search on sentence
#             emb = model.encode([sent], normalize_embeddings=True)
#             D, I = index.search(emb, top_k)
#             for idx in I[0]:
#                 if str(idx) in meta:
#                     skills.append(meta[str(idx)])

#     # Deduplicate
#     seen = set()
#     unique_skills = []
#     for s in skills:
#         if s["uri"] not in seen:
#             unique_skills.append(s)
#             seen.add(s["uri"])

#     return unique_skills


# # ---------- Match Occupations ----------
# def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
#     skill_uris = {s["uri"] for s in skills}
#     occ_matches = []

#     for _, row in occ_df.iterrows():
#         occ_skill_uris = set(row["skillUris"])
#         overlap = skill_uris.intersection(occ_skill_uris)
#         if overlap:
#             matched_labels = [uri_to_meta[uri]["preferredLabel"] for uri in overlap if uri in uri_to_meta]
#             occ_matches.append({
#                 "preferredLabel": row["preferredLabel"],
#                 "conceptUri": row["conceptUri"],
#                 "iscoGroup": {
#                     "code": row["code"],
#                     "name": row["preferredLabel_isco"],
#                     "description": row["description"]
#                 },
#                 "matching_skills_count": len(overlap),
#                 "matched_skills": matched_labels,
#                 "required_skills_count": len(occ_skill_uris)
#             })

#     occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
#     return occ_matches[:top_k]


# # ---------- Rank skills by embedding similarity ----------
# def rank_skills_by_similarity(entries, resume_text, model, reference_embs=None, top_n=15):
#     if not entries:
#         return []

#     labels = [e["preferredLabel"] for e in entries]
#     embs = model.encode(labels, normalize_embeddings=True)

#     # reference: either candidate embeddings or resume embedding
#     if reference_embs is None:
#         reference_embs = model.encode([resume_text], normalize_embeddings=True)

#     sims = []
#     for i, emb in enumerate(embs):
#         score = float(np.max(emb @ reference_embs.T))
#         sims.append((entries[i], score))

#     return [
#         {
#             "preferredLabel": e["preferredLabel"],
#             "uri": e["uri"],
#             "description": e.get("description", ""),
#             "similarity": sim
#         }
#         for e, sim in sorted(sims, key=lambda x: x[1], reverse=True)[:top_n]
#     ]


# # ---------- Step 5: Skill Gap Analysis with Synonyms + Ranking ----------
# def skill_gap_summary(top_skills, chosen_role, uri_to_meta, resume_text, model):
#     student_labels = [s["preferredLabel"] for s in top_skills]
#     student_embs = model.encode(student_labels, normalize_embeddings=True)

#     required_uris = set(chosen_role.get("skillUris", []))
#     required_entries = [uri_to_meta[u] for u in required_uris if u in uri_to_meta]

#     strengths, gaps = [], []

#     for entry in required_entries:
#         label = entry.get("preferredLabel", "")
#         synonyms = [label]

#         alt_labels = entry.get("altLabels", "")
#         if isinstance(alt_labels, str) and alt_labels:
#             synonyms.extend([x.strip() for x in alt_labels.split(",") if x.strip()])

#         if any(s.lower() in [x.lower() for x in student_labels] for s in synonyms):
#             strengths.append(entry)
#         else:
#             gaps.append(entry)

#     return {
#         "strengths": rank_skills_by_similarity(strengths, resume_text, model, student_embs, top_n=20),
#         "gaps": rank_skills_by_similarity(gaps, resume_text, model, student_embs, top_n=30),
#         "underdeveloped": []
#     }


# # ---------- Main Pipeline ----------
# def parse_resume(file_path: str, project_root: str):
#     print("⚙️ Loading CareerBERT model...")
#     model = SentenceTransformer("lwolfrum2/careerbert-jg")

#     print("📄 Reading resume...")
#     resume_text = load_resume_text(file_path)

#     print("📦 Loading ESCO index...")
#     index, meta, uri_to_meta = load_esco_index(project_root)

#     print("👔 Loading occupations + ISCO groups...")
#     occ_df = load_occupations(project_root)

#     print("🔍 Extracting skills with hybrid method...")
#     top_skills = extract_top_skills(resume_text, model, index, meta, top_k=5)

#     print("🧭 Matching occupations...")
#     top_occupations = match_occupations(top_skills, occ_df, uri_to_meta, top_k=10)

#     # Step 5: Gap analysis on best role (if available)
#     chosen_role_uri = top_occupations[0]["conceptUri"] if top_occupations else None
#     chosen_role_row = occ_df.loc[occ_df["conceptUri"] == chosen_role_uri].iloc[0] if chosen_role_uri else None
#     skill_gap = skill_gap_summary(
#         top_skills,
#         chosen_role_row if chosen_role_row is not None else {},
#         uri_to_meta,
#         resume_text,
#         model
#     )

#     results = {
#         "top_skills_identified": top_skills[:15],
#         "top_matching_occupations": top_occupations[:10],
#         "skill_gap_analysis": skill_gap
#     }

#     out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(results, f, indent=4, ensure_ascii=False)

#     print(f"✅ Results saved to {out_path}")


# # ---------- Run ----------
# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Usage: python app/parse_resume.py <resume.pdf/docx>")
#         sys.exit(1)

#     project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
#     file_path = sys.argv[1]
#     parse_resume(file_path, project_root)




import os
import sys
import json
import faiss
import pandas as pd
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
import spacy
from spacy.matcher import PhraseMatcher
from rapidfuzz import process  # for fuzzy matching
import re
import numpy as np
from docx import Document   # for .docx resumes


# ---------- Load Resume Text ----------
def load_resume_text(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        reader = PdfReader(file_path)
        text = " ".join([page.extract_text() or "" for page in reader.pages])
    elif file_path.lower().endswith(".docx"):
        doc = Document(file_path)
        text = " ".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Please use .pdf or .docx")
    return text.strip()


# ---------- Load ESCO Skill Index ----------
def load_esco_index(project_root: str):
    vector_store_dir = os.path.join(project_root, "vector_store")
    index_path = os.path.join(vector_store_dir, "esco_faiss.index")
    meta_path = os.path.join(vector_store_dir, "esco_meta.json")

    if not os.path.exists(index_path) or not os.path.exists(meta_path):
        raise FileNotFoundError("❌ Missing ESCO FAISS index or metadata. Run build_esco_index.py first.")

    index = faiss.read_index(index_path)
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    # Reverse lookup: conceptUri → metadata
    uri_to_meta = {v["uri"]: v for v in meta.values() if "uri" in v}
    return index, meta, uri_to_meta


# ---------- Load Occupations + ISCO + Relations ----------
def load_occupations(project_root: str):
    occ_path = os.path.join(project_root, "data", "esco", "occupations_en.csv")
    isco_path = os.path.join(project_root, "data", "esco", "ISCOGroups_en.csv")
    relations_path = os.path.join(project_root, "data", "esco", "occupationSkillRelations_en.csv")

    occ_df = pd.read_csv(occ_path, dtype=str).fillna("")
    isco_df = pd.read_csv(isco_path, dtype=str).fillna("")
    rel_df = pd.read_csv(relations_path, dtype=str).fillna("")

    # Merge ISCO group info into occupations
    occ_df = occ_df.merge(
        isco_df[["code", "preferredLabel", "description"]],
        left_on="iscoGroup",
        right_on="code",
        how="left",
        suffixes=("", "_isco"),
    )

    # Build skillUris list per occupation
    occ_skills = rel_df.groupby("occupationUri")["skillUri"].apply(list).to_dict()
    occ_df["skillUris"] = occ_df["conceptUri"].map(occ_skills).apply(lambda x: x if isinstance(x, list) else [])

    return occ_df


# ---------- Detect Candidate Skills (NER + dictionary) ----------
def detect_skills(resume_text: str, meta):
    nlp = spacy.load("en_core_web_md")
    matcher = PhraseMatcher(nlp.vocab, attr="LOWER")

    # collect ESCO labels + synonyms
    all_labels = []
    for entry in meta.values():
        all_labels.append(entry.get("preferredLabel", ""))
        alt_labels = entry.get("altLabels", "")
        if isinstance(alt_labels, str) and alt_labels:
            all_labels.extend([x.strip() for x in alt_labels.split(",")])

    patterns = [nlp.make_doc(label) for label in set(all_labels) if label]
    matcher.add("SKILLS", patterns)

    doc = nlp(resume_text)
    matches = matcher(doc)

    found = [doc[start:end].text for _, start, end in matches]
    return list(set(found))


# ---------- Normalize Mentions with Synonyms/Fuzzy Matching ----------
def normalize_skill(skill, meta):
    labels_map = {v["preferredLabel"]: v for v in meta.values() if "preferredLabel" in v}
    match, score, _ = process.extractOne(skill, labels_map.keys())
    if match and score > 80:
        return labels_map[match]
    return None

import re
import numpy as np
import spacy
from spacy.matcher import PhraseMatcher
from rapidfuzz import process

# Make sure spaCy is loaded once
nlp = spacy.load("en_core_web_md")

STOPWORDS = set([
    "the","and","or","of","in","on","to","for","with","by","at","an","a","is","as","be","are","was","were","it",
    "use","using","used"
])

def word_overlap(text1, text2):
    """Compute word overlap ratio ignoring stopwords and very short tokens."""
    t1 = {w.lower() for w in re.findall(r"\w+", text1) if len(w) > 2 and w.lower() not in STOPWORDS}
    t2 = {w.lower() for w in re.findall(r"\w+", text2) if len(w) > 2 and w.lower() not in STOPWORDS}
    return len(t1.intersection(t2)) > 0

def extract_top_skills(resume_text: str, model, index, meta, top_k=15):
    """
    Hybrid skill extraction with overlap-aware filtering:
    - PhraseMatcher + fuzzy detection
    - Embedding fallback with stricter cutoff and word overlap check
    - Weighted by section (Skills > Experience > Others)
    - Deduplicated and ranked
    """

    # --- Section splitting & weighting ---
    section_weights = {}
    sections = re.split(r'\n(?=[A-Z][A-Z ]{2,}:?)', resume_text)
    for sec in sections:
        sec_name = sec.strip().split("\n", 1)[0].lower()
        weight = 1.0
        if "skill" in sec_name:
            weight = 3.0
        elif "experience" in sec_name or "project" in sec_name:
            weight = 2.0
        section_weights[sec] = weight

    skills = []

    for sec, weight in section_weights.items():
        sentences = re.split(r'(?<=[.!?])\s+', sec)

        for sent in sentences:
            if not sent.strip():
                continue

            sent_emb = model.encode([sent], normalize_embeddings=True)

            # --- Phrase/fuzzy detection ---
            candidate_mentions = detect_skills(sent, meta)
            if candidate_mentions:
                for mention in candidate_mentions:
                    norm_entry = normalize_skill(mention, meta)
                    query_label = norm_entry["preferredLabel"] if norm_entry else mention
                    emb = model.encode([query_label], normalize_embeddings=True)
                    D, I = index.search(emb, 5)

                    for idx in I[0]:
                        if str(idx) in meta:
                            sim = np.dot(emb, sent_emb.T).item()
                            entry = meta[str(idx)].copy()
                            entry["detection_source"] = "phrase_or_fuzzy"
                            entry["similarity"] = sim * weight
                            skills.append(entry)
            else:
                # --- Embedding fallback (stricter rules) ---
                D, I = index.search(sent_emb, 5)
                for idx in I[0]:
                    if str(idx) in meta:
                        entry = meta[str(idx)].copy()
                        label = entry["preferredLabel"]

                        # Compute sim vs sentence
                        emb = model.encode([label], normalize_embeddings=True)
                        sim = np.dot(emb, sent_emb.T).item()

                        # Apply stricter cutoff + overlap
                        if sim >= 0.75 and word_overlap(label, sent):
                            entry["detection_source"] = "embedding_fallback"
                            entry["similarity"] = sim * weight
                            skills.append(entry)

    # --- Deduplication (keep max sim per URI) ---
    dedup = {}
    for s in skills:
        if s["uri"] not in dedup or s["similarity"] > dedup[s["uri"]]["similarity"]:
            dedup[s["uri"]] = s

    ranked = sorted(dedup.values(), key=lambda x: x["similarity"], reverse=True)

    return ranked[:top_k]





# ---------- Match Occupations ----------
def match_occupations(skills, occ_df, uri_to_meta, top_k=5):
    skill_uris = {s["uri"] for s in skills}
    occ_matches = []

    for _, row in occ_df.iterrows():
        occ_skill_uris = set(row["skillUris"])
        overlap = skill_uris.intersection(occ_skill_uris)
        if overlap:
            matched_labels = [uri_to_meta[uri]["preferredLabel"] for uri in overlap if uri in uri_to_meta]
            occ_matches.append({
                "preferredLabel": row["preferredLabel"],
                "conceptUri": row["conceptUri"],
                "iscoGroup": {
                    "code": row["code"],
                    "name": row["preferredLabel_isco"],
                    "description": row["description"]
                },
                "matching_skills_count": len(overlap),
                "matched_skills": matched_labels,
                "required_skills_count": len(occ_skill_uris)
            })

    occ_matches = sorted(occ_matches, key=lambda x: x["matching_skills_count"], reverse=True)
    return occ_matches[:top_k]


# ---------- Step 5: Skill Gap Analysis ----------
def skill_gap_summary(top_skills, chosen_role, uri_to_meta, resume_text, model):
    student_labels = [s["preferredLabel"] for s in top_skills]
    student_embs = model.encode(student_labels, normalize_embeddings=True) if student_labels else None

    required_uris = set(chosen_role.get("skillUris", []))
    required_entries = [uri_to_meta[u] for u in required_uris if u in uri_to_meta]

    strengths, gaps = [], []

    for entry in required_entries:
        label = entry.get("preferredLabel", "")
        synonyms = [label]

        alt_labels = entry.get("altLabels", "")
        if isinstance(alt_labels, str) and alt_labels:
            synonyms.extend([x.strip() for x in alt_labels.split(",") if x.strip()])

        if any(s.lower() in [x.lower() for x in student_labels] for s in synonyms):
            strengths.append(entry)
        else:
            gaps.append(entry)

    def rank(entries, ref_embs):
        if not entries:
            return []
        labels = [e["preferredLabel"] for e in entries]
        embs = model.encode(labels, normalize_embeddings=True)
        sims = []
        for i, emb in enumerate(embs):
            score = float(np.max(emb @ ref_embs.T)) if ref_embs is not None else 0.0
            sims.append((entries[i], score))
        return [
            {
                "preferredLabel": e["preferredLabel"],
                "uri": e["uri"],
                "description": e.get("description", ""),
                "similarity": sim
            }
            for e, sim in sorted(sims, key=lambda x: x[1], reverse=True)
        ]

    return {
        "strengths": rank(strengths, student_embs)[:20],
        "gaps": rank(gaps, student_embs)[:30],
        "underdeveloped": []
    }


# ---------- Main Pipeline ----------
def parse_resume(file_path: str, project_root: str):
    print("⚙️ Loading CareerBERT model...")
    model = SentenceTransformer("lwolfrum2/careerbert-jg")

    print("📄 Reading resume...")
    resume_text = load_resume_text(file_path)

    print("📦 Loading ESCO index...")
    index, meta, uri_to_meta = load_esco_index(project_root)

    print("👔 Loading occupations + ISCO groups...")
    occ_df = load_occupations(project_root)

    print("🔍 Extracting skills with hybrid method...")
    top_skills = extract_top_skills(resume_text, model, index, meta, top_k=5)

    print("🧭 Matching occupations...")
    top_occupations = match_occupations(top_skills, occ_df, uri_to_meta, top_k=10)

    chosen_role_uri = top_occupations[0]["conceptUri"] if top_occupations else None
    chosen_role_row = occ_df.loc[occ_df["conceptUri"] == chosen_role_uri].iloc[0] if chosen_role_uri else None
    skill_gap = skill_gap_summary(
        top_skills,
        chosen_role_row if chosen_role_row is not None else {},
        uri_to_meta,
        resume_text,
        model
    )

    results = {
        "top_skills_identified": top_skills[:15],
        "top_matching_occupations": top_occupations[:10],
        "skill_gap_analysis": skill_gap
    }

    out_path = os.path.join(project_root, "data", "resumes", "sample_skills.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=4, ensure_ascii=False)

    print(f"✅ Results saved to {out_path}")


# ---------- Run ----------
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python app/parse_resume.py <resume.pdf/docx>")
        sys.exit(1)

    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    file_path = sys.argv[1]
    parse_resume(file_path, project_root)

